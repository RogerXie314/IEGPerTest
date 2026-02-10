import json
import os
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

from docx import Document


@dataclass
class TableRef:
    tableIndex: int
    rows: List[List[str]]


def _norm_cell_text(s: str) -> str:
    return (s or "").replace("\r", "\n").strip()


def read_docx(docx_path: str) -> Document:
    return Document(docx_path)


def extract_paragraphs(doc: Document) -> List[str]:
    paras: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paras.append(t)
    return paras


def extract_tables(doc: Document) -> List[List[List[str]]]:
    tables: List[List[List[str]]] = []
    for t in doc.tables:
        rows: List[List[str]] = []
        for row in t.rows:
            rows.append([_norm_cell_text(c.text) for c in row.cells])
        tables.append(rows)
    return tables


def main() -> None:
    workspace = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    docx_path = os.path.join(workspace, "docs", "IEG&EDR V300R011C01&C11 接口设计说明书.docx")
    if not os.path.exists(docx_path):
        raise FileNotFoundError(docx_path)

    doc = read_docx(docx_path)
    paras = extract_paragraphs(doc)
    tables = extract_tables(doc)

    endpoint_re = re.compile(r"/USM/[A-Za-z0-9_]+\.do")

    # Endpoint contexts from paragraphs
    endpoints: Dict[str, List[str]] = {}

    def add_endpoint(ep: str, ctx: str) -> None:
        if ep not in endpoints:
            endpoints[ep] = []
        # de-dup but preserve order
        if ctx not in endpoints[ep]:
            endpoints[ep].append(ctx)

    for i, text in enumerate(paras):
        found = endpoint_re.findall(text)
        if not found:
            continue
        ctx_parts = []
        if i > 0:
            ctx_parts.append(paras[i - 1])
        ctx_parts.append(text)
        if i + 1 < len(paras):
            ctx_parts.append(paras[i + 1])
        ctx = "\n".join(ctx_parts)
        for ep in found:
            add_endpoint(ep, ctx)

    # Tables that mention endpoint
    endpoint_tables: Dict[str, List[Dict[str, Any]]] = {}
    for ti, rows in enumerate(tables):
        flat = "\n".join("\t".join(r) for r in rows)
        found = endpoint_re.findall(flat)
        if not found:
            continue
        for ep in sorted(set(found)):
            endpoint_tables.setdefault(ep, []).append({"tableIndex": ti, "rows": rows})

    # Heuristic: rows that likely talk about cmd id/type
    cmd_keywords_re = re.compile(r"CMDID|CmdId|cmdid|CMDTYPE|CmdType|cmdtype|命令字|命令ID|指令ID|指令字")
    number_re = re.compile(r"\b\d{2,4}\b")
    cmd_rows: List[Dict[str, Any]] = []

    def consider_row(ti: int, ri: int, row: List[str]) -> None:
        joined = " ".join([c for c in row if c])
        if not joined:
            return
        if not cmd_keywords_re.search(joined):
            return
        nums = number_re.findall(joined)
        if not nums:
            return
        cmd_rows.append({"tableIndex": ti, "rowIndex": ri, "row": row, "numbers": nums})

    for ti, rows in enumerate(tables):
        for ri, row in enumerate(rows):
            consider_row(ti, ri, row)

    out: Dict[str, Any] = {
        "source": os.path.relpath(docx_path, workspace).replace("\\", "/"),
        "paragraphCount": len(paras),
        "tableCount": len(tables),
        "endpoints": endpoints,
        "endpointTables": endpoint_tables,
        "cmdIdRows": cmd_rows,
    }

    artifacts_dir = os.path.join(workspace, "artifacts")
    os.makedirs(artifacts_dir, exist_ok=True)
    out_path = os.path.join(artifacts_dir, "interface_doc_index.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)

    print("Wrote", os.path.relpath(out_path, workspace).replace("\\", "/"))
    print("Endpoints found:", len(endpoints))
    for ep in sorted(endpoints.keys())[:50]:
        print(" ", ep)


if __name__ == "__main__":
    main()
